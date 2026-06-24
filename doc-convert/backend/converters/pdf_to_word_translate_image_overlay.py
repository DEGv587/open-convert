"""
PDF 转 Word（带翻译）- 图片叠加模式
使用百度图片识别 API 获取坐标，在原图上叠加翻译文字
"""
import os
from typing import Optional, Callable, List, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO

import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

from translators.baidu_translate import get_translator


def convert(
    input_path: str,
    output_path: str,
    progress_cb: Optional[Callable] = None,
    target_lang: str = 'zh',
    source_lang: str = 'auto',
    start_page: Optional[int] = None,
    end_page: Optional[int] = None
):
    """
    PDF 转 Word（带翻译图片叠加）

    Args:
        input_path: 输入 PDF 路径
        output_path: 输出 DOCX 路径
        progress_cb: 进度回调 progress_cb(percent, stage)
        target_lang: 目标语言（'zh' | 'en'）
        source_lang: 源语言（忽略）
        start_page: 起始页码（从 0 开始）
        end_page: 结束页码（从 0 开始）
    """
    translator = get_translator()

    # 打开 PDF
    doc_pdf = fitz.open(input_path)
    total_pages = len(doc_pdf)

    # 计算实际处理的页码范围
    actual_start = start_page or 0
    actual_end = end_page if end_page is not None else total_pages - 1
    actual_end = min(actual_end, total_pages - 1)
    page_range = range(actual_start, actual_end + 1)

    if progress_cb:
        progress_cb(5, f"使用百度图片识别 API 处理第 {actual_start + 1}-{actual_end + 1} 页")

    # 步骤1：百度图片识别 + 图片叠加（10-90%）
    translated_images = _process_pages_with_overlay(
        doc_pdf, progress_cb, page_range, translator, target_lang
    )

    doc_pdf.close()

    if progress_cb:
        progress_cb(90, f"创建 Word 文档")

    # 步骤2：创建 Word 文档（90-100%）
    doc = Document()

    # 设置页面尺寸（A4）
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # 插入翻译后的图片
    for idx, img_bytes in enumerate(translated_images):
        if not img_bytes:
            continue

        page_num = page_range.start + idx + 1

        # 添加页码标记（可选）
        page_marker = doc.add_paragraph()
        page_marker_run = page_marker.add_run(f"第 {page_num} 页")
        page_marker_run.font.size = Inches(0.12)
        page_marker.alignment = 1  # 居中

        # 插入图片（占满页面宽度）
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(BytesIO(img_bytes), width=Inches(7.5))

        # 分页
        if idx < len(translated_images) - 1:
            doc.add_page_break()

        if progress_cb:
            pct = 90 + int((idx + 1) / len(translated_images) * 10)
            progress_cb(pct, f"写入文档 {idx + 1}/{len(translated_images)} 页")

    # 保存
    doc.save(output_path)

    if progress_cb:
        progress_cb(100, f"完成（已翻译 {len(translated_images)} 页）")


def _process_pages_with_overlay(
    doc_pdf,
    progress_cb: Optional[Callable],
    page_range: range,
    translator,
    target_lang: str
) -> List[bytes]:
    """
    使用百度图片识别 API 处理页面并叠加翻译

    Returns:
        List[bytes]: 每页叠加翻译后的图片（PNG 格式）
    """
    page_count = len(page_range)
    processed_images = [None] * page_count
    completed = 0

    # 确定翻译方向
    from_lang = 'en' if target_lang == 'zh' else 'zh'
    to_lang = target_lang

    def process_page(idx: int, page_num: int) -> Tuple[int, Optional[bytes]]:
        """处理单页：识别 + 叠加翻译"""
        page = doc_pdf[page_num]

        # 渲染页面为图片（先尝试 150 DPI，如果超过 4MB 或不符合百度限制则降低）
        # 百度限制：最短边 >= 30px，最长边 <= 4096px，长宽比 3:1 以内，大小 <= 4MB
        dpi = 150
        while dpi >= 72:
            pix = page.get_pixmap(dpi=dpi)
            img_bytes = pix.tobytes("png")

            # 获取图片尺寸
            width, height = pix.width, pix.height
            min_side = min(width, height)
            max_side = max(width, height)
            aspect_ratio = max_side / min_side if min_side > 0 else 0
            img_size_mb = len(img_bytes) / (1024 * 1024)

            # 检查是否符合百度限制
            if (img_size_mb <= 3.8 and  # 留一点余量
                min_side >= 30 and
                max_side <= 4096 and
                aspect_ratio <= 3.0):
                break

            # 不符合，降低 DPI
            dpi -= 25
            if dpi < 72:
                # 最低 72 DPI 仍不符合，强制使用并记录警告
                dpi = 72
                pix = page.get_pixmap(dpi=dpi)
                img_bytes = pix.tobytes("png")
                print(f"[警告] 页 {page_num + 1} 尺寸: {pix.width}x{pix.height}, 大小: {len(img_bytes)/1024/1024:.1f}MB")
                break

        try:
            # 调用百度图片识别 API
            result = translator.recognize_image(img_bytes, from_lang=from_lang, to_lang=to_lang)

            if not result or not result.get('blocks'):
                # 识别失败，返回原图
                print(f"[页 {page_num + 1}] 百度图片识别无结果")
                return idx, img_bytes

            # 打开图片准备叠加（需要重新从字节加载，因为前面的 pix 是 PyMuPDF 对象）
            img = Image.open(BytesIO(img_bytes)).convert('RGB')
            draw = ImageDraw.Draw(img)

            # 加载中文字体（更大字号，确保清晰）
            font_size = 40  # 从 28 提升到 40
            try:
                # macOS 中文字体
                font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", font_size)
            except:
                try:
                    # macOS 备用：华文黑体
                    font = ImageFont.truetype("/System/Library/Fonts/STHeiti Light.ttc", font_size)
                except:
                    try:
                        # Linux 中文字体
                        font = ImageFont.truetype("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", font_size)
                    except:
                        try:
                            # Linux 备用
                            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
                        except:
                            # 降级到默认字体
                            font = ImageFont.load_default()

            # 遍历每个文本块，叠加翻译
            for block in result['blocks']:
                dst = block.get('dst', '').strip()
                if not dst:
                    continue

                # 解析坐标 "x y w h"
                rect_str = block.get('rect', '')
                if not rect_str:
                    continue

                try:
                    x, y, w, h = map(int, rect_str.split())
                except:
                    continue

                # 绘制半透明背景（白色）
                overlay = Image.new('RGBA', img.size, (255, 255, 255, 0))
                overlay_draw = ImageDraw.Draw(overlay)
                overlay_draw.rectangle(
                    [(x, y), (x + w, y + h)],
                    fill=(255, 255, 255, 220)  # 半透明白色
                )
                img = Image.alpha_composite(img.convert('RGBA'), overlay).convert('RGB')
                draw = ImageDraw.Draw(img)

                # 绘制翻译文字（自动换行）
                _draw_multiline_text(draw, dst, (x, y, w, h), font)

            # 转换为 PNG 字节
            output = BytesIO()
            img.save(output, format='PNG')
            return idx, output.getvalue()

        except Exception as e:
            print(f"[页 {page_num + 1}] 处理失败: {e}")
            return idx, img_bytes  # 返回原图

    # 并行处理（控制并发数）
    max_workers = min(4, page_count)

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_page, idx, page_num): idx
                   for idx, page_num in enumerate(page_range)}

        for future in as_completed(futures):
            idx, img_bytes = future.result()
            processed_images[idx] = img_bytes
            completed += 1

            if progress_cb and completed % 2 == 0:
                pct = 10 + int(completed / page_count * 80)
                progress_cb(pct, f"图片识别+叠加翻译 {completed}/{page_count} 页")

    return processed_images


def _draw_multiline_text(draw, text: str, rect: Tuple[int, int, int, int], font):
    """在指定矩形区域内绘制多行文字（自动换行）"""
    x, y, w, h = rect
    max_width = w - 10  # 留边距

    # 简单换行策略：按字符宽度估算
    lines = []
    current_line = ""

    for char in text:
        test_line = current_line + char
        # 估算宽度（PIL 的 textbbox 需要单独计算）
        try:
            bbox = draw.textbbox((0, 0), test_line, font=font)
            text_width = bbox[2] - bbox[0]
        except:
            text_width = len(test_line) * 15  # 降级估算

        if text_width <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = char

    if current_line:
        lines.append(current_line)

    # 绘制每行
    line_height = 32
    for i, line in enumerate(lines):
        draw.text(
            (x + 5, y + 5 + i * line_height),
            line,
            fill=(50, 50, 50),  # 深灰色
            font=font
        )
