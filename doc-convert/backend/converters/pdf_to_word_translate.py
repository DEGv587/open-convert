"""
PDF 转 Word（带翻译）
方案4：基于 pdf2docx 保留原文档格式，在段落后插入翻译
支持并行处理和 OCR 识别
"""
import os
from typing import Optional, Callable, List, Tuple
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from translators.baidu_translate import get_translator


def convert(
    input_path: str,
    output_path: str,
    progress_cb: Optional[Callable] = None,
    target_lang: str = 'zh',
    source_lang: str = 'auto',
    start_page: Optional[int] = None,
    end_page: Optional[int] = None
):
    """
    PDF 转 Word（带翻译，保留原文档格式）

    Args:
        input_path: 输入 PDF 路径
        output_path: 输出 DOCX 路径
        progress_cb: 进度回调 progress_cb(percent, stage)
        target_lang: 目标语言（'zh' | 'en'）
        source_lang: 源语言（'auto' 自动检测）
        start_page: 起始页码（从 0 开始，None 表示从第一页）
        end_page: 结束页码（从 0 开始，None 表示到最后一页）
    """
    translator = get_translator()

    # 先检测是否需要 OCR
    doc_pdf = fitz.open(input_path)
    total_pages = len(doc_pdf)

    # 计算实际处理的页码范围
    actual_start = start_page or 0
    actual_end = end_page if end_page is not None else total_pages - 1
    actual_end = min(actual_end, total_pages - 1)

    page_range = range(actual_start, actual_end + 1)
    page_count = len(page_range)

    # 检测是否需要 OCR（采样指定范围的前10页）
    need_ocr = _detect_need_ocr(doc_pdf, sample_pages=min(10, page_count), start_page=actual_start)

    if need_ocr:
        # 扫描版 PDF：直接创建 Word，OCR + 翻译写入
        if progress_cb:
            progress_cb(5, f"检测到扫描版 PDF，OCR 识别第 {actual_start + 1}-{actual_end + 1} 页")

        _convert_scanned_pdf_with_translation(
            doc_pdf, output_path, translator, progress_cb,
            page_range, target_lang, source_lang
        )
        doc_pdf.close()
        return  # 扫描版 PDF 处理完成，直接返回

    # 文本版 PDF：使用 pdf2docx 保留格式
    doc_pdf.close()

    # 步骤1：pdf2docx 转换（保留格式，10%进度）
    if progress_cb:
        progress_cb(5, "使用 pdf2docx 转换（保留格式）")

    cv = Converter(input_path)
    cv.convert(output_path, start=start_page or 0, end=end_page)
    cv.close()

    if progress_cb:
        progress_cb(10, "格式转换完成，准备提取文本")

    # 重新打开用于文本提取
    doc_pdf = fitz.open(input_path)

    # 步骤2：提取文本（10-40%）
    if progress_cb:
        progress_cb(10, f"提取文本第 {actual_start + 1}-{actual_end + 1} 页")

    page_texts = []
    for idx, page_num in enumerate(page_range):
        page = doc_pdf[page_num]
        text = page.get_text()
        page_texts.append(text)

        if progress_cb and idx % 10 == 0:
            pct = 10 + int((idx + 1) / page_count * 30)
            progress_cb(pct, f"提取文本 {idx + 1}/{page_count} 页")

    doc_pdf.close()

    # 合并所有文本并按段落分割
    all_text = '\n\n'.join(page_texts)
    paragraphs = [p.strip() for p in all_text.split('\n\n') if p.strip() and len(p.strip()) > 5]

    if not paragraphs:
        raise ValueError("PDF 中未提取到有效文本")

    if progress_cb:
        progress_cb(40, f"提取到 {len(paragraphs)} 个段落，准备翻译")

    # 步骤3：并行翻译（40-70%）
    if target_lang == 'zh':
        from_lang = 'en' if source_lang == 'auto' else source_lang
        to_lang = 'zh'
    else:
        from_lang = 'zh' if source_lang == 'auto' else source_lang
        to_lang = 'en'

    def translate_progress(current, total):
        if progress_cb:
            pct = 40 + int(current / total * 30)
            progress_cb(pct, f"翻译中 {current}/{total} 段")

    translation_result = translator.translate_batch(
        paragraphs,
        from_lang=from_lang,
        to_lang=to_lang,
        progress_callback=translate_progress
    )

    # 提取翻译结果
    translated_paragraphs = translation_result['results']
    success_count = translation_result['success_count']
    total_count = translation_result['total']
    quota_exceeded = translation_result['quota_exceeded']

    # 如果限额用完，给用户警告（但继续处理已翻译的部分）
    if quota_exceeded:
        warning_msg = f"翻译 API 限额已用完，仅完成 {success_count}/{total_count} 段翻译"
        if progress_cb:
            progress_cb(70, warning_msg)
    else:
        if progress_cb:
            progress_cb(70, f"翻译完成（{success_count}/{total_count}），插入译文到文档")

    # 步骤4：打开转换后的 Word，在段落后插入翻译（70-100%）
    doc_word = Document(output_path)

    # 构建原文→译文映射
    translation_map = dict(zip(paragraphs, translated_paragraphs))

    # 收集需要插入的位置（先收集再批量插入，避免索引问题）
    insertions = []  # [(index, translation), ...]

    for i, para in enumerate(doc_word.paragraphs):
        para_text = para.text.strip()

        # 跳过空段落
        if not para_text or len(para_text) < 5:
            continue

        # 尝试匹配翻译
        matched_translation = None
        for original, translated in translation_map.items():
            if _text_similarity(original, para_text) > 0.7:  # 相似度阈值
                matched_translation = translated
                break

        if matched_translation:
            insertions.append((i, matched_translation))

    if progress_cb:
        progress_cb(75, f"匹配到 {len(insertions)} 段，插入翻译")

    # 批量插入（从后往前，避免索引偏移）
    for idx, (para_idx, translation) in enumerate(reversed(insertions)):
        _insert_translation_after_paragraph(doc_word, para_idx, translation)

        if progress_cb and idx % 10 == 0:
            pct = 75 + int(idx / len(insertions) * 25)
            progress_cb(pct, f"插入译文 {idx}/{len(insertions)}")

    # 如果限额用完，在文档开头添加警告
    if quota_exceeded:
        warning_para = doc_word.paragraphs[0].insert_paragraph_before(
            f"⚠️ 翻译提示：百度翻译 API 限额已用完，仅完成 {success_count}/{total_count} 段翻译。"
            f"未翻译的段落保留原文。"
        )
        run = warning_para.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(200, 0, 0)
        run.bold = True
        warning_para.paragraph_format.space_after = Pt(12)

    # 保存前统一字体（避免字体缺失警告）
    _normalize_fonts(doc_word)

    # 保存
    doc_word.save(output_path)

    # 返回时带上限额信息
    if progress_cb:
        if quota_exceeded:
            progress_cb(100, f"完成（部分翻译 {success_count}/{total_count}，API 限额已用完）")
        else:
            progress_cb(100, f"完成（插入 {len(insertions)} 段翻译）")


def _detect_need_ocr(doc_pdf, sample_pages: int = 10, start_page: int = 0) -> bool:
    """检测是否需要 OCR（采样前N页）"""
    total_chars = 0
    end_sample = min(start_page + sample_pages, len(doc_pdf))
    for i in range(start_page, end_sample):
        text = doc_pdf[i].get_text()
        total_chars += len(text.strip())

    avg_chars_per_page = total_chars / (end_sample - start_page) if end_sample > start_page else 0
    # 如果平均每页少于50个字符，认为是扫描版
    return avg_chars_per_page < 50


def _baidu_image_recognize_pages(
    doc_pdf,
    progress_cb: Optional[Callable],
    page_range: range,
    translator,
    target_lang: str
) -> List[Tuple[str, str]]:
    """
    使用百度图片识别 API 处理页面（OCR + 翻译一步完成）

    Args:
        doc_pdf: PDF 文档对象
        progress_cb: 进度回调
        page_range: 页码范围
        translator: 百度翻译实例
        target_lang: 目标语言（'zh' | 'en'）

    Returns:
        List[(original_text, translated_text)] 每页的原文和译文
    """
    page_count = len(page_range)
    page_results = [('', '')] * page_count
    completed = 0

    # 确定翻译方向
    from_lang = 'en' if target_lang == 'zh' else 'zh'
    to_lang = target_lang

    def process_page(idx: int, page_num: int) -> Tuple[int, str, str]:
        """处理单页：提取图片 + 百度识别"""
        page = doc_pdf[page_num]

        # 渲染页面为图片
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")

        try:
            # 调用百度图片识别 API
            result = translator.recognize_image(img_bytes, from_lang=from_lang, to_lang=to_lang)

            if result:
                original = result.get('original_text', '')
                translated = result.get('translated_text', '')
                return idx, original, translated
            else:
                return idx, '', ''

        except Exception as e:
            print(f"[页 {page_num}] 百度图片识别失败: {e}")
            return idx, '', ''

    # 并行处理（控制并发数，避免超过 QPS 限制）
    max_workers = min(4, page_count)  # 百度图片识别 QPS 限制较低

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_page, idx, page_num): idx
                   for idx, page_num in enumerate(page_range)}

        for future in as_completed(futures):
            idx, original, translated = future.result()
            page_results[idx] = (original, translated)
            completed += 1

            if progress_cb and completed % 2 == 0:
                pct = 10 + int(completed / page_count * 60)
                progress_cb(pct, f"图片识别+翻译 {completed}/{page_count} 页")

    return page_results


def _text_similarity(text1: str, text2: str) -> float:
    """计算两段文本的相似度（简单的字符串匹配）"""
    if not text1 or not text2:
        return 0.0

    # 归一化：去除多余空格和换行
    t1 = ' '.join(text1.split())
    t2 = ' '.join(text2.split())

    # 完全匹配
    if t1 == t2:
        return 1.0

    # 包含关系
    if t1 in t2 or t2 in t1:
        shorter = min(len(t1), len(t2))
        longer = max(len(t1), len(t2))
        return shorter / longer

    # 前缀匹配（前50个字符）
    prefix_len = min(50, len(t1), len(t2))
    if t1[:prefix_len] == t2[:prefix_len]:
        return 0.8

    return 0.0


def _insert_translation_after_paragraph(doc: Document, para_index: int, translation: str):
    """在指定段落后插入翻译（灰色小字）"""
    # python-docx 没有直接的 insert_after，用变通方法
    para = doc.paragraphs[para_index]

    # 在下一个段落前插入，如果是最后一段则直接添加
    if para_index + 1 < len(doc.paragraphs):
        # 获取下一个段落的 XML 元素
        next_para = doc.paragraphs[para_index + 1]
        new_para_element = next_para._element.getparent().makeelement(
            qn('w:p'),
            nsmap=next_para._element.nsmap
        )
        next_para._element.addprevious(new_para_element)

        # 创建段落对象
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_para_element, doc)

        # 添加文本
        run = new_para.add_run(translation)
    else:
        # 最后一段，直接添加
        new_para = doc.add_paragraph()
        run = new_para.add_run(translation)

    # 设置样式：灰色、小字、斜体
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    # 段落格式
    new_para.paragraph_format.left_indent = Pt(12)
    new_para.paragraph_format.space_after = Pt(6)


def _normalize_fonts(doc):
    """统一文档字体，避免字体缺失警告"""
    # 常见的跨平台字体
    # 中文：宋体/SimSun (Windows), 华文宋体/STSong (macOS)
    # 英文：Arial, Calibri
    
    fallback_fonts = {
        'zh': 'SimSun',      # 宋体（Windows/Linux 都支持）
        'en': 'Arial'        # Arial（跨平台）
    }
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                # 检测是否包含中文
                has_chinese = any('一' <= c <= '鿿' for c in run.text)
                # 设置合适的后备字体
                run.font.name = fallback_fonts['zh'] if has_chinese else fallback_fonts['en']
    
    # 处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name:
                            has_chinese = any('一' <= c <= '鿿' for c in run.text)
                            run.font.name = fallback_fonts['zh'] if has_chinese else fallback_fonts['en']


def _convert_scanned_pdf_with_translation(
    doc_pdf,
    output_path: str,
    translator,
    progress_cb: Optional[Callable],
    page_range: range,
    target_lang: str,
    source_lang: str
):
    """扫描版 PDF 专用：OCR + 翻译直接写入 Word"""
    
    page_count = len(page_range)

    # 步骤1：并行 OCR（5-40%）
    page_texts = _parallel_ocr_extract(doc_pdf, progress_cb, page_range=page_range, translator=translator)

    # 按页分组段落，并过滤重复内容
    all_paragraphs_by_page = []
    for page_idx, page_text in enumerate(page_texts):
        page_paras = [p.strip() for p in page_text.split('\n\n')
                     if p.strip() and len(p.strip()) > 5 and not _is_gibberish(p.strip())]
        all_paragraphs_by_page.append({
            'page_num': page_range.start + page_idx + 1,  # 实际页码（1-indexed）
            'paragraphs': page_paras
        })

    # 收集所有段落用于过滤
    all_paragraphs = []
    for page_data in all_paragraphs_by_page:
        all_paragraphs.extend(page_data['paragraphs'])

    if not all_paragraphs:
        raise ValueError("OCR 未提取到有效文本")

    original_count = len(all_paragraphs)

    # 过滤重复内容并保持页面分组
    filtered_set = set(_filter_repetitive_content(all_paragraphs))

    # 应用过滤到每页
    for page_data in all_paragraphs_by_page:
        page_data['paragraphs'] = [p for p in page_data['paragraphs'] if p in filtered_set]

    # 重新收集过滤后的段落
    all_paragraphs = []
    for page_data in all_paragraphs_by_page:
        all_paragraphs.extend(page_data['paragraphs'])

    if progress_cb:
        filtered_info = f"（过滤 {original_count - len(all_paragraphs)} 个重复段落）" if len(all_paragraphs) < original_count else ""
        progress_cb(40, f"提取到 {len(all_paragraphs)} 个段落{filtered_info}，准备翻译")
    
    # 步骤2：翻译（40-70%）
    if target_lang == 'zh':
        from_lang = 'en' if source_lang == 'auto' else source_lang
        to_lang = 'zh'
    else:
        from_lang = 'zh' if source_lang == 'auto' else source_lang
        to_lang = 'en'
    
    def translate_progress(current, total):
        if progress_cb:
            pct = 40 + int(current / total * 30)
            progress_cb(pct, f"翻译中 {current}/{total} 段")
    
    translation_result = translator.translate_batch(
        all_paragraphs,
        from_lang=from_lang,
        to_lang=to_lang,
        progress_callback=translate_progress
    )

    translated_paragraphs = translation_result['results']
    success_count = translation_result['success_count']
    total_count = translation_result['total']
    quota_exceeded = translation_result['quota_exceeded']
    
    if quota_exceeded:
        warning_msg = f"翻译 API 限额已用完，仅完成 {success_count}/{total_count} 段翻译"
        if progress_cb:
            progress_cb(70, warning_msg)
    else:
        if progress_cb:
            progress_cb(70, f"翻译完成（{success_count}/{total_count}），创建文档")
    
    # 步骤3：创建 Word 文档（70-100%）
    doc = Document()

    # 如果限额用完，添加警告
    if quota_exceeded:
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run(
            f"⚠️ 翻译提示：百度翻译 API 限额已用完，仅完成 {success_count}/{total_count} 段翻译。"
            f"未翻译的段落保留原文。"
        )
        warning_run.font.size = Pt(10)
        warning_run.font.color.rgb = RGBColor(200, 0, 0)
        warning_run.bold = True
        warning_para.paragraph_format.space_after = Pt(12)

    # 创建翻译映射
    translation_map = dict(zip(all_paragraphs, translated_paragraphs))

    # 按页写入，每页前添加页码分隔
    written_count = 0
    for page_data in all_paragraphs_by_page:
        page_num = page_data['page_num']
        page_paras = page_data['paragraphs']

        if not page_paras:
            continue

        # 添加页码分隔（居中）
        page_marker = doc.add_paragraph()
        page_marker_run = page_marker.add_run(f"—————  第 {page_num} 页  —————")
        page_marker_run.font.size = Pt(10)
        page_marker_run.font.color.rgb = RGBColor(150, 150, 150)
        page_marker_run.bold = True
        page_marker.alignment = 1  # 居中对齐
        page_marker.paragraph_format.space_after = Pt(24)  # 页码后空两行

        # 写入该页的原文 + 翻译
        for original in page_paras:
            translated = translation_map.get(original, original)

            # 原文段落
            para_orig = doc.add_paragraph()
            run_orig = para_orig.add_run(original)
            run_orig.font.name = 'Arial'
            run_orig.font.size = Pt(11)
            para_orig.paragraph_format.space_after = Pt(6)

            # 翻译段落（灰色小字斜体）
            para_trans = doc.add_paragraph()
            run_trans = para_trans.add_run(translated)
            run_trans.font.name = 'SimSun'
            run_trans.font.size = Pt(9)
            run_trans.font.color.rgb = RGBColor(100, 100, 100)
            run_trans.italic = True
            para_trans.paragraph_format.left_indent = Pt(12)
            para_trans.paragraph_format.space_after = Pt(12)

            written_count += 1
            if progress_cb and written_count % 10 == 0:
                pct = 70 + int(written_count / len(all_paragraphs) * 30)
                progress_cb(pct, f"写入文档 {written_count}/{len(all_paragraphs)} 段")

        # 页面结束后额外空一行
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 保存
    doc.save(output_path)
    
    if progress_cb:
        if quota_exceeded:
            progress_cb(100, f"完成（部分翻译 {success_count}/{total_count}，API 限额已用完）")
        else:
            progress_cb(100, f"完成（已翻译 {len(all_paragraphs)} 段）")


def _filter_repetitive_content(paragraphs: List[str]) -> List[str]:
    """
    过滤重复的页脚、版权信息等

    策略：
    1. 识别常见的页脚模式（版权、单元标题等）
    2. 如果某段落与其他段落高度相似且出现多次，只保留第一次
    """
    if len(paragraphs) < 5:
        return paragraphs

    # 定义页脚关键词模式
    footer_keywords = [
        'publishers',
        'copyright',
        '©',
        'pte ltd',
        'unit 1',
        'unit 2',
        'unit 3',
        'unit 4',
        'unit 5',
        'learning mathematics book'
    ]

    def is_likely_footer(text: str) -> bool:
        """判断是否可能是页脚"""
        text_lower = text.lower()
        # 短文本 + 包含关键词
        return len(text) < 150 and any(kw in text_lower for kw in footer_keywords)

    def normalize_for_comparison(text: str) -> str:
        """标准化文本用于比较"""
        import re
        # 去除所有空白字符，转小写
        text = re.sub(r'\s+', '', text.lower())
        # 去除所有标点和特殊字符（包括 ©）
        text = re.sub(r'[^\w]', '', text)
        return text

    # 收集可能的页脚
    potential_footers = []
    for i, para in enumerate(paragraphs):
        if is_likely_footer(para):
            potential_footers.append((i, para, normalize_for_comparison(para)))

    if not potential_footers:
        return paragraphs

    # 找出重复的页脚（标准化后相同）
    from collections import Counter
    normalized_texts = [norm for _, _, norm in potential_footers]
    counter = Counter(normalized_texts)

    # 如果某个标准化文本出现 >= 2次，认为是重复页脚
    repetitive_normalized = {text for text, count in counter.items() if count >= 2}

    if not repetitive_normalized:
        return paragraphs

    # 标记要删除的索引（保留第一次出现）
    seen_normalized = set()
    indices_to_remove = set()

    for idx, original, normalized in potential_footers:
        if normalized in repetitive_normalized:
            if normalized in seen_normalized:
                # 后续重复，标记删除
                indices_to_remove.add(idx)
            else:
                # 第一次出现，保留
                seen_normalized.add(normalized)

    # 过滤
    filtered = [para for i, para in enumerate(paragraphs) if i not in indices_to_remove]

    removed_count = len(paragraphs) - len(filtered)
    if removed_count > 0:
        print(f"[过滤] 移除 {removed_count} 个重复页脚段落")

    return filtered


def _is_gibberish(text: str) -> bool:
    """判断是否是 OCR 乱码"""
    import re

    if len(text) < 3:
        return False

    # 规则1：包含明显的乱码模式
    gibberish_patterns = [
        r'MAOA+D',  # MAOAOAAD
        r'DODD\s*D',  # DODD D
        r'GOC+O+d',  # GOCCOOOOd
        r'[A-Z]{2,}OA',  # 连续大写字母+OA
    ]
    for pattern in gibberish_patterns:
        if re.search(pattern, text):
            return True

    # 规则2：连续的大写字母/数字组合且无意义（超过5个连续大写）
    if re.search(r'[A-Z]{5,}', text) and not re.search(r'\b[A-Z]{2,5}\b', text):
        return True

    # 规则3：大量重复字符（单个字符占比超过40%）
    for char in set(text):
        if char.isalnum() and text.count(char) > len(text) * 0.4:
            return True

    # 规则4：包含方括号且内容是大写字母组合（如 [ MAOAOAAD）
    if re.search(r'\[\s*[A-Z\s]+$', text):
        return True

    return False
