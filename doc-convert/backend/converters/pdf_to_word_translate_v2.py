"""
PDF 转 Word（带翻译）- 百度图片识别版
方案：使用百度图片识别 API（OCR + 翻译一步完成），避免 Tesseract 乱码
"""
import os
from typing import Optional, Callable, List, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor

from translators.baidu_translate import get_translator


def convert(
    input_path: str,
    output_path: str,
    progress_cb: Optional[Callable] = None,
    target_lang: str = 'zh',
    source_lang: str = 'auto',
    start_page: Optional[int] = None,
    end_page: Optional[int] = None
):
    """
    PDF 转 Word（带翻译，使用百度图片识别 API）

    Args:
        input_path: 输入 PDF 路径
        output_path: 输出 DOCX 路径
        progress_cb: 进度回调 progress_cb(percent, stage)
        target_lang: 目标语言（'zh' | 'en'）
        source_lang: 源语言（忽略，API 自动检测）
        start_page: 起始页码（从 0 开始）
        end_page: 结束页码（从 0 开始）
    """
    translator = get_translator()

    # 打开 PDF
    doc_pdf = fitz.open(input_path)
    total_pages = len(doc_pdf)

    # 计算实际处理的页码范围
    actual_start = start_page or 0
    actual_end = end_page if end_page is not None else total_pages - 1
    actual_end = min(actual_end, total_pages - 1)
    page_range = range(actual_start, actual_end + 1)

    if progress_cb:
        progress_cb(5, f"使用百度图片识别 API 处理第 {actual_start + 1}-{actual_end + 1} 页")

    # 步骤1：百度图片识别（10-70%）
    page_results = _baidu_image_recognize_pages(
        doc_pdf, progress_cb, page_range, translator, target_lang
    )

    doc_pdf.close()

    # 步骤2：按页组织内容（过滤空白和乱码）
    all_paragraphs_by_page = []

    for idx, (original_text, translated_text) in enumerate(page_results):
        page_num = page_range.start + idx + 1

        # 百度 API 返回的文本已经是按行分隔的（每个 block 一行）
        # 直接按换行符分割即可
        original_lines = [line.strip() for line in original_text.split('\n') if line.strip()]
        translated_lines = [line.strip() for line in translated_text.split('\n') if line.strip()]

        # 确保原文和译文数量一致（配对）
        min_len = min(len(original_lines), len(translated_lines))
        original_lines = original_lines[:min_len]
        translated_lines = translated_lines[:min_len]

        # 配对存储
        pairs = list(zip(original_lines, translated_lines))

        all_paragraphs_by_page.append({
            'page_num': page_num,
            'pairs': pairs  # [(original, translated), ...]
        })

    # 收集所有原文用于过滤重复
    all_originals = []
    for page_data in all_paragraphs_by_page:
        all_originals.extend([p[0] for p in page_data['pairs']])

    original_count = len(all_originals)

    if not all_originals:
        raise ValueError("图片识别未提取到有效文本")

    # 过滤重复内容（页脚等）- 只过滤长度较短且重复的文本块
    filtered_originals_set = set(_filter_repetitive_content(all_originals))

    # 应用过滤到每页（保留译文配对）
    for page_data in all_paragraphs_by_page:
        page_data['pairs'] = [(o, t) for o, t in page_data['pairs']
                              if o in filtered_originals_set]

    # 统计
    total_pairs = sum(len(p['pairs']) for p in all_paragraphs_by_page)

    if progress_cb:
        filtered_info = f"（过滤 {original_count - total_pairs} 个重复段落）" if total_pairs < original_count else ""
        progress_cb(70, f"提取到 {total_pairs} 个段落{filtered_info}，创建文档")

    # 步骤3：创建 Word 文档（70-100%）
    doc = Document()

    # 按页写入
    written_count = 0
    for page_data in all_paragraphs_by_page:
        page_num = page_data['page_num']
        pairs = page_data['pairs']

        if not pairs:
            continue

        # 添加页码分隔（居中）
        page_marker = doc.add_paragraph()
        page_marker_run = page_marker.add_run(f"—————  第 {page_num} 页  —————")
        page_marker_run.font.size = Pt(10)
        page_marker_run.font.color.rgb = RGBColor(150, 150, 150)
        page_marker_run.bold = True
        page_marker.alignment = 1  # 居中对齐
        page_marker.paragraph_format.space_after = Pt(24)  # 页码后空两行

        # 写入原文 + 翻译
        for original, translated in pairs:
            # 跳过过短的文本（可能是页码或标点）
            if len(original.strip()) < 3:
                continue

            # 原文段落
            para_orig = doc.add_paragraph()
            run_orig = para_orig.add_run(original)
            run_orig.font.name = 'Arial'
            run_orig.font.size = Pt(11)
            para_orig.paragraph_format.space_after = Pt(6)

            # 翻译段落（灰色小字斜体）
            para_trans = doc.add_paragraph()
            run_trans = para_trans.add_run(translated)
            run_trans.font.name = 'SimSun'
            run_trans.font.size = Pt(9)
            run_trans.font.color.rgb = RGBColor(100, 100, 100)
            run_trans.italic = True
            para_trans.paragraph_format.left_indent = Pt(12)
            para_trans.paragraph_format.space_after = Pt(12)

            written_count += 1
            if progress_cb and written_count % 10 == 0:
                pct = 70 + int(written_count / total_pairs * 30)
                progress_cb(pct, f"写入文档 {written_count}/{total_pairs} 段")

        # 页面结束后额外空一行
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 保存
    doc.save(output_path)

    if progress_cb:
        progress_cb(100, f"完成（已翻译 {total_pairs} 段）")


def _baidu_image_recognize_pages(
    doc_pdf,
    progress_cb: Optional[Callable],
    page_range: range,
    translator,
    target_lang: str
) -> List[Tuple[str, str]]:
    """
    使用百度图片识别 API 处理页面（OCR + 翻译一步完成）

    Returns:
        List[(original_text, translated_text)] 每页的原文和译文
    """
    page_count = len(page_range)
    page_results = [('', '')] * page_count
    completed = 0

    # 确定翻译方向
    from_lang = 'en' if target_lang == 'zh' else 'zh'
    to_lang = target_lang

    def process_page(idx: int, page_num: int) -> Tuple[int, str, str]:
        """处理单页：提取图片 + 百度识别"""
        page = doc_pdf[page_num]

        # 渲染页面为图片
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")

        try:
            # 调用百度图片识别 API
            result = translator.recognize_image(img_bytes, from_lang=from_lang, to_lang=to_lang)

            if result:
                original = result.get('original_text', '')
                translated = result.get('translated_text', '')
                return idx, original, translated
            else:
                return idx, '', ''

        except Exception as e:
            print(f"[页 {page_num + 1}] 百度图片识别失败: {e}")
            return idx, '', ''

    # 并行处理（控制并发数，避免超过 QPS 限制）
    max_workers = min(4, page_count)  # 百度图片识别 QPS 限制较低

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_page, idx, page_num): idx
                   for idx, page_num in enumerate(page_range)}

        for future in as_completed(futures):
            idx, original, translated = future.result()
            page_results[idx] = (original, translated)
            completed += 1

            if progress_cb and completed % 2 == 0:
                pct = 10 + int(completed / page_count * 60)
                progress_cb(pct, f"图片识别+翻译 {completed}/{page_count} 页")

    return page_results


def _filter_repetitive_content(paragraphs: List[str]) -> List[str]:
    """
    过滤重复的页脚、版权信息等

    策略：
    1. 识别常见的页脚模式（版权、单元标题等）
    2. 如果某段落与其他段落高度相似且出现多次，只保留第一次
    """
    if len(paragraphs) < 5:
        return paragraphs

    # 定义页脚关键词模式
    footer_keywords = [
        'publishers',
        'copyright',
        '©',
        'pte ltd',
        'unit 1',
        'unit 2',
        'unit 3',
        'unit 4',
        'unit 5',
        'learning mathematics book'
    ]

    def is_likely_footer(text: str) -> bool:
        """判断是否可能是页脚"""
        text_lower = text.lower()
        # 短文本 + 包含关键词
        return len(text) < 150 and any(kw in text_lower for kw in footer_keywords)

    def normalize_for_comparison(text: str) -> str:
        """标准化文本用于比较"""
        import re
        # 去除所有空白字符，转小写
        text = re.sub(r'\s+', '', text.lower())
        # 去除所有标点和特殊字符（包括 ©）
        text = re.sub(r'[^\w]', '', text)
        return text

    # 收集可能的页脚
    potential_footers = []
    for i, para in enumerate(paragraphs):
        if is_likely_footer(para):
            potential_footers.append((i, para, normalize_for_comparison(para)))

    if not potential_footers:
        return paragraphs

    # 找出重复的页脚（标准化后相同）
    from collections import Counter
    normalized_texts = [norm for _, _, norm in potential_footers]
    counter = Counter(normalized_texts)

    # 如果某个标准化文本出现 >= 2次，认为是重复页脚
    repetitive_normalized = {text for text, count in counter.items() if count >= 2}

    if not repetitive_normalized:
        return paragraphs

    # 标记要删除的索引（保留第一次出现）
    seen_normalized = set()
    indices_to_remove = set()

    for idx, original, normalized in potential_footers:
        if normalized in repetitive_normalized:
            if normalized in seen_normalized:
                # 后续重复，标记删除
                indices_to_remove.add(idx)
            else:
                # 第一次出现，保留
                seen_normalized.add(normalized)

    # 过滤
    filtered = [para for i, para in enumerate(paragraphs) if i not in indices_to_remove]

    removed_count = len(paragraphs) - len(filtered)
    if removed_count > 0:
        print(f"[过滤] 移除 {removed_count} 个重复页脚段落")

    return filtered
