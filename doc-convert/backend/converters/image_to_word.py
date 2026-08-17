"""OCR images into an editable Word document."""

import os
import re
from pathlib import Path
from typing import Sequence, Union

import pytesseract
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image, ImageOps, UnidentifiedImageError
from pytesseract import Output
from pytesseract.pytesseract import TesseractError, TesseractNotFoundError


OCR_LANG = os.getenv("OCR_LANG", "chi_sim+eng")
OCR_CONFIG = os.getenv("OCR_CONFIG", "--oem 1 --psm 3")
OCR_TIMEOUT_SECONDS = int(os.getenv("OCR_TIMEOUT_SECONDS", "90"))
_CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
_NO_SPACE_BEFORE = set(",.!?;:，。！？；：、)]}》」』）］】〉")
_NO_SPACE_AFTER = set("([{<《「『（［【〈")


def _as_paths(input_paths: Union[str, Sequence[str]]) -> list[str]:
    if isinstance(input_paths, (str, Path)):
        paths = [str(input_paths)]
    else:
        paths = [str(path) for path in input_paths]
    if not paths:
        raise ValueError("至少需要一张图片")
    return paths


def _load_for_ocr(path: str) -> Image.Image:
    """Load an image with camera orientation applied and improve small scans."""
    try:
        with Image.open(path) as source:
            source.load()
            image = ImageOps.exif_transpose(source)
            if image.mode in ("RGBA", "LA") or "transparency" in image.info:
                rgba = image.convert("RGBA")
                white = Image.new("RGBA", rgba.size, "white")
                image = Image.alpha_composite(white, rgba)
            image = image.convert("RGB")
    except (UnidentifiedImageError, OSError) as exc:
        raise ValueError(f"无法读取图片：{Path(path).name}") from exc

    # Upscaling small screenshots gives Tesseract more pixels for Chinese glyphs.
    longest_edge = max(image.size)
    if longest_edge < 1600:
        scale = min(2.0, 1600 / longest_edge)
        image = image.resize(
            (round(image.width * scale), round(image.height * scale)),
            Image.Resampling.LANCZOS,
        )
    return ImageOps.autocontrast(ImageOps.grayscale(image))


def _check_ocr_runtime() -> None:
    try:
        available = set(pytesseract.get_languages(config=""))
    except TesseractNotFoundError as exc:
        raise RuntimeError("服务器未安装 Tesseract OCR") from exc
    except TesseractError as exc:
        raise RuntimeError(f"Tesseract OCR 初始化失败：{exc}") from exc

    requested = {lang for lang in OCR_LANG.split("+") if lang}
    missing = sorted(requested - available)
    if missing:
        raise RuntimeError(
            f"服务器缺少 OCR 语言包：{', '.join(missing)}，请安装 chi_sim 和 eng 语言包"
        )


def _is_cjk(text: str) -> bool:
    return bool(text) and bool(_CJK_RE.search(text))


def _join_tokens(tokens: list[str]) -> str:
    """Join OCR words without inserting spaces inside Chinese text."""
    result = ""
    for raw_token in tokens:
        token = re.sub(r"\s+", " ", raw_token).strip()
        if not token:
            continue
        if not result:
            result = token
            continue

        previous = result[-1]
        first = token[0]
        previous_cjk = _is_cjk(previous)
        first_cjk = _is_cjk(first)
        needs_space = (
            previous not in _NO_SPACE_AFTER
            and first not in _NO_SPACE_BEFORE
            and (
                (previous_cjk == first_cjk and not previous_cjk)
                or (previous_cjk != first_cjk and previous.isalnum() and first.isalnum())
            )
        )
        result += (" " if needs_space else "") + token
    return result.strip()


def _extract_paragraphs(image: Image.Image) -> list[list[str]]:
    """Return OCR lines grouped into reading-order paragraphs."""
    try:
        data = pytesseract.image_to_data(
            image,
            lang=OCR_LANG,
            config=OCR_CONFIG,
            output_type=Output.DICT,
            timeout=OCR_TIMEOUT_SECONDS,
        )
    except TesseractNotFoundError as exc:
        raise RuntimeError("服务器未安装 Tesseract OCR") from exc
    except TesseractError as exc:
        raise RuntimeError(f"图片 OCR 失败：{exc}") from exc

    lines: dict[tuple[int, int, int], dict] = {}
    count = len(data.get("text", []))
    for index in range(count):
        text = re.sub(r"\s+", " ", str(data["text"][index])).strip()
        if not text:
            continue
        try:
            block = int(data["block_num"][index])
            paragraph = int(data["par_num"][index])
            line = int(data["line_num"][index])
            left = int(data["left"][index])
            top = int(data["top"][index])
        except (KeyError, TypeError, ValueError):
            continue

        key = (block, paragraph, line)
        item = lines.setdefault(key, {"tokens": [], "left": left, "top": top})
        item["tokens"].append((left, text))
        item["left"] = min(item["left"], left)
        item["top"] = min(item["top"], top)

    # Tesseract's block numbers are not guaranteed to be visual reading order.
    ordered_lines = []
    for key, item in lines.items():
        tokens = [text for _, text in sorted(item["tokens"], key=lambda pair: pair[0])]
        joined = _join_tokens(tokens)
        if joined:
            ordered_lines.append({**item, "key": key, "text": joined})
    ordered_lines.sort(key=lambda item: (item["top"], item["left"]))

    paragraphs: list[list[str]] = []
    paragraph_keys: dict[tuple[int, int], list[dict]] = {}
    for item in ordered_lines:
        # Group lines by Tesseract paragraph, then sort the groups visually.
        paragraph_keys.setdefault(item["key"][:2], []).append(item)

    groups = sorted(
        paragraph_keys.values(),
        key=lambda group: (min(item["top"] for item in group), min(item["left"] for item in group)),
    )
    for group in groups:
        paragraphs.append([item["text"] for item in sorted(group, key=lambda item: (item["top"], item["left"]))])
    return paragraphs


def _set_run_font(run) -> None:
    """Set both Latin and East Asian font slots in the DOCX XML."""
    run.font.name = "Arial"
    run.font.size = Pt(11)
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        from docx.oxml import OxmlElement
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "宋体")


def _append_paragraph(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _set_run_font(run)


def convert(input_paths: Union[str, Sequence[str]], output_path: str, progress_cb=None):
    """OCR each image in order and write editable Unicode text to DOCX."""
    paths = _as_paths(input_paths)
    _check_ocr_runtime()

    document = Document()
    document.core_properties.title = Path(paths[0]).stem
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if progress_cb:
        progress_cb(5, "准备 OCR 识别...")

    for index, path in enumerate(paths):
        if index:
            document.add_page_break()
        if progress_cb:
            progress_cb(10 + round(index / len(paths) * 70), f"正在识别第 {index + 1}/{len(paths)} 张图片...")

        paragraphs = _extract_paragraphs(_load_for_ocr(path))
        for lines in paragraphs:
            _append_paragraph(document, lines)

    if progress_cb:
        progress_cb(85, "正在生成可编辑 Word...")
    document.save(output_path)
    if progress_cb:
        progress_cb(100, "Word 文档已生成")
