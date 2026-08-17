"""Convert pasted Markdown/LaTeX text to an editable DOCX."""

from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from converters.latex2omml import normalize, to_omml


CJK_FONT = "宋体"
CJK_FONT_BOLD = "黑体"
LATIN_FONT = "Times New Roman"
HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
ITEM_RE = re.compile(r"^(\d+)\.\s*(.*)$")
INLINE_MATH_RE = re.compile(r"\$(?!\$)(.+?)(?<!\$)\$|\\\((.+?)\\\)", re.S)
ANSWER_ROW_RE = re.compile(r"^(?:\d+\s*[.．]\s*[A-Da-d]\s*[　\s]*){3,}$")
SECTION_HEADING_RE = re.compile(r"^[一二三四五六七八九十]+[、.．]\s*.+")
TITLE_HINT_RE = re.compile(r"(?:试卷|测试题?|练习题?|考试|答案|报告|文档|方案|通知|总结)$")
BARE_DISPLAY_START_RE = re.compile(r"^\\begin\{(?:align\*?|aligned|gather\*?)\}")
BARE_DISPLAY_END_RE = re.compile(r"\\end\{(?:align\*?|aligned|gather\*?)\}\s*$")
IMPLICIT_MATH_RE = re.compile(
    r"[A-Za-z0-9\\{}^_()\[\]+*/=<>.,:;!|%&\-\s‐‑–—−×÷·≤≥≠≈]+"
)
OPTION_PREFIX_RE = re.compile(r"(?<![A-Za-z0-9])([A-Da-d])\.\s*")
MATH_COMMAND_RE = re.compile(
    r"\\(?:d?frac|tfrac|sqrt|boldsymbol|mathbf|bf|mathrm|text|operatorname|"
    r"cdot|times|div|pm|mp|neq|ne|leq|le|geq|ge|approx|infty|pi|"
    r"alpha|beta|gamma|delta|theta|lambda|mu|sigma|angle|triangle|circ)(?![A-Za-z])"
)
BLANK_LINE_RE = re.compile(r"(?:[＿﹍﹎﹏‗]\s*){3,}|(?:_\s*){3,}")


def _set_fonts(run, bold: bool = False) -> None:
    run.font.name = LATIN_FONT
    run.font.size = Pt(10.5)
    run_pr = run._element.get_or_add_rPr()
    rfonts = run_pr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run_pr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT_BOLD if bold else CJK_FONT)


def _append_text_run(paragraph, text: str, bold: bool = False, underline: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(normalize(text))
    run.bold = bold
    run.underline = underline
    _set_fonts(run, bold)


def _add_text(paragraph, text: str, bold: bool = False) -> None:
    position = 0
    for match in BLANK_LINE_RE.finditer(text):
        _append_text_run(paragraph, text[position:match.start()], bold)
        marker = match.group(0)
        marker_count = sum(char in "＿﹍﹎﹏‗_" for char in marker)
        _append_text_run(paragraph, "_" * marker_count, bold)
        position = match.end()
    _append_text_run(paragraph, text[position:], bold)


def _add_math(paragraph, latex: str, math_mode: str = "omml") -> None:
    try:
        if math_mode == "image":
            # Imported lazily so normal editable Word conversion does not pay
            # Matplotlib's startup and font-cache cost.
            from converters.latex_to_image import render_latex_png

            run = paragraph.add_run()
            run.add_picture(BytesIO(render_latex_png(latex)))
            # Inline drawings sit slightly high next to Chinese body text.
            # Lowering by one point keeps the mathematical baseline aligned.
            run_pr = run._element.get_or_add_rPr()
            position = OxmlElement("w:position")
            position.set(qn("w:val"), "-2")
            run_pr.append(position)
        else:
            paragraph._p.append(parse_xml(to_omml(latex)))
    except Exception:
        # A malformed formula should remain visible and editable as source text.
        _add_text(paragraph, latex)


def _looks_like_implicit_math(candidate: str) -> bool:
    # Normalize typography copied from browsers/PDFs before detection.  In
    # particular, U+2212 (−) must remain inside a LaTeX command's braces; if it
    # splits ``\boldsymbol{...}``, the unmatched half is emitted as raw source.
    value = normalize(candidate).strip()
    if not value:
        return False
    if MATH_COMMAND_RE.search(value) or "^" in value or "_" in value:
        return True
    if re.search(r"[A-Za-z0-9)]\s*[=+*/<>-]\s*[(A-Za-z0-9-]", value):
        return True
    return bool(re.search(r"(?:\d+[A-Za-z]|[A-Za-z]\d+)", value))


def _add_math_candidate(paragraph, candidate: str, bold: bool = False, math_mode: str = "omml") -> None:
    parts = re.split(r"(?=(?<![A-Za-z0-9])[A-Da-d]\.\s*)", candidate)
    for part in parts:
        if not part:
            continue
        option = re.match(r"^([A-Da-d]\.)(\s*)(.*)$", part, re.S)
        if option:
            body = option.group(3).rstrip(" ")
            trailing = option.group(3)[len(option.group(3).rstrip(" ")):]
            _add_text(paragraph, option.group(1) + option.group(2), bold)
            if body:
                _add_math(paragraph, body, math_mode)
            _add_text(paragraph, trailing, bold)
        else:
            _add_math(paragraph, part.strip(), math_mode)


def _add_candidate_with_blanks(
    paragraph, candidate: str, bold: bool = False, math_mode: str = "omml"
) -> None:
    position = 0
    for match in BLANK_LINE_RE.finditer(candidate):
        math_part = candidate[position:match.start()]
        if math_part:
            if _looks_like_implicit_math(math_part):
                _add_math_candidate(paragraph, math_part, bold, math_mode)
            else:
                _add_text(paragraph, math_part, bold)
        _add_text(paragraph, match.group(0), bold)
        position = match.end()
    remainder = candidate[position:]
    if remainder:
        if _looks_like_implicit_math(remainder):
            _add_math_candidate(paragraph, remainder, bold, math_mode)
        else:
            _add_text(paragraph, remainder, bold)


def _add_implicit_math(paragraph, text: str, bold: bool = False, math_mode: str = "omml") -> None:
    position = 0
    for match in IMPLICIT_MATH_RE.finditer(text):
        candidate = match.group(0)
        is_option_math = bool(
            OPTION_PREFIX_RE.search(candidate)
            and re.search(r"(?:\\|\^|_|\d|[=+*/<>-])", candidate)
        )
        if not _looks_like_implicit_math(candidate) and not is_option_math:
            continue
        _add_text(paragraph, text[position:match.start()], bold)
        leading = candidate[:len(candidate) - len(candidate.lstrip(" "))]
        trailing = candidate[len(candidate.rstrip(" ")):]
        _add_text(paragraph, leading, bold)
        _add_candidate_with_blanks(paragraph, candidate.strip(), bold, math_mode)
        _add_text(paragraph, trailing, bold)
        position = match.end()
    _add_text(paragraph, text[position:], bold)


def _add_rich(
    paragraph,
    text: str,
    bold: bool = False,
    detect_implicit: bool = True,
    math_mode: str = "omml",
) -> None:
    position = 0
    for match in INLINE_MATH_RE.finditer(text):
        plain = text[position:match.start()]
        if detect_implicit:
            _add_implicit_math(paragraph, plain, bold, math_mode)
        else:
            _add_text(paragraph, plain, bold)
        _add_math(paragraph, match.group(1) or match.group(2), math_mode)
        position = match.end()
    plain = text[position:]
    if detect_implicit:
        _add_implicit_math(paragraph, plain, bold, math_mode)
    else:
        _add_text(paragraph, plain, bold)


def _paragraph(document, *, indent=0.0, hanging=0.0, align=None, space_after=4.0, space_before=0.0):
    paragraph = document.add_paragraph()
    fmt = paragraph.paragraph_format
    if indent:
        fmt.left_indent = Cm(indent)
    if hanging:
        fmt.first_line_indent = Cm(-hanging)
    if align is not None:
        fmt.alignment = align
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = 1.35
    return paragraph


def _setup_page(document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    style = document.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def _split_display_blocks(lines: list[str]) -> list[tuple[str, str]]:
    blocks = []
    end_mode = None
    buffer = []
    for line in lines:
        stripped = line.strip()
        if end_mode is None and stripped.startswith("$$"):
            rest = stripped[2:]
            if rest.endswith("$$") and len(rest) >= 2:
                blocks.append(("display", rest[:-2]))
            else:
                end_mode = "dollars"
                buffer = [rest] if rest else []
            continue
        if end_mode is None and BARE_DISPLAY_START_RE.match(stripped):
            buffer = [stripped]
            if BARE_DISPLAY_END_RE.search(stripped):
                blocks.append(("display", stripped))
                buffer = []
            else:
                end_mode = "environment"
            continue
        if end_mode is not None:
            if end_mode == "dollars" and stripped.endswith("$$"):
                buffer.append(stripped[:-2])
                blocks.append(("display", "\n".join(buffer)))
                end_mode = None
                buffer = []
            elif end_mode == "environment" and BARE_DISPLAY_END_RE.search(stripped):
                buffer.append(stripped)
                blocks.append(("display", "\n".join(buffer)))
                end_mode = None
                buffer = []
            else:
                buffer.append(line)
            continue
        blocks.append(("text", line))
    if end_mode is not None and buffer:
        blocks.append(("display", "\n".join(buffer)))
    return blocks


def _add_title(document, text: str, first: bool) -> None:
    if not first:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    paragraph = _paragraph(document, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    run = paragraph.add_run(normalize(text))
    run.bold = True
    run.font.size = Pt(18)
    _set_fonts(run, bold=True)


def _add_heading(document, text: str) -> None:
    paragraph = _paragraph(document, space_before=10, space_after=6)
    run = paragraph.add_run(normalize(text))
    run.bold = True
    run.font.size = Pt(12)
    _set_fonts(run, bold=True)


def convert(text: str, output_path: str, progress_cb=None, *, math_mode: str = "omml"):
    if not isinstance(text, str) or not text.strip():
        raise ValueError("粘贴的文本不能为空")
    if math_mode not in {"omml", "image"}:
        raise ValueError(f"不支持的公式渲染模式: {math_mode}")
    if progress_cb:
        progress_cb(10, "正在解析文本和数学公式...")

    document = Document()
    _setup_page(document)
    seen_title = False
    seen_content = False
    in_item = False
    blocks = _split_display_blocks(text.splitlines())

    for kind, payload in blocks:
        if kind == "display":
            paragraph = _paragraph(document, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=6)
            _add_math(paragraph, payload, math_mode)
            continue

        line = payload.rstrip()
        if not line.strip():
            in_item = False
            continue
        heading = HEADING_RE.match(line)
        if heading:
            if len(heading.group(1)) == 1:
                _add_title(document, heading.group(2), first=not seen_title)
                seen_title = True
            else:
                _add_heading(document, heading.group(2))
            in_item = False
            seen_content = True
            continue
        stripped = line.strip()
        if stripped == "参考答案":
            _add_title(document, stripped, first=not seen_title)
            seen_title = True
            seen_content = True
            in_item = False
            continue
        if not seen_content and len(stripped) <= 60 and TITLE_HINT_RE.search(stripped):
            _add_title(document, stripped, first=True)
            seen_title = True
            seen_content = True
            in_item = False
            continue
        if SECTION_HEADING_RE.match(stripped):
            _add_heading(document, stripped)
            seen_content = True
            in_item = False
            continue
        if ANSWER_ROW_RE.match(stripped):
            paragraph = _paragraph(document)
            _add_rich(paragraph, stripped, detect_implicit=False, math_mode=math_mode)
            in_item = False
            seen_content = True
            continue
        item = ITEM_RE.match(line)
        if item:
            paragraph = _paragraph(document, indent=0.85, hanging=0.85)
            _add_text(paragraph, f"{item.group(1)}. ")
            _add_rich(paragraph, item.group(2), math_mode=math_mode)
            in_item = True
            seen_content = True
            continue
        paragraph = _paragraph(document, indent=0.85 if in_item else 0, space_after=2 if in_item else 4)
        _add_rich(paragraph, stripped, math_mode=math_mode)
        seen_content = True

    if progress_cb:
        progress_cb(85, "正在生成 PDF 排版..." if math_mode == "image" else "正在生成 Word 文档...")
    document.save(output_path)
    if progress_cb:
        progress_cb(100, "PDF 排版已生成" if math_mode == "image" else "Word 文档已生成")
