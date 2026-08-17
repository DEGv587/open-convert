import tempfile
import unittest
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from PIL import Image

from converters.text_to_word import convert


class TextToWordTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.work_dir = Path(self.temp_dir.name)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_latex_is_written_as_native_omml(self):
        source = (
            "# 公式测试\n\n"
            "1. 计算 $x^2 + \\dfrac{1}{2} + \\sqrt{y}$\n\n"
            "$$\\begin{align*}\n"
            "y^2&=x^2+1\\\\\n"
            "z&=\\sqrt{y}\n"
            "\\end{align*}$$"
        )
        output = self.work_dir / "formula.docx"

        convert(source, str(output))

        document = Document(output)
        self.assertIn("公式测试", "\n".join(paragraph.text for paragraph in document.paragraphs))
        with zipfile.ZipFile(output) as package:
            xml = package.read("word/document.xml").decode("utf-8")
        self.assertGreaterEqual(xml.count("<m:oMath"), 2)
        self.assertIn("<m:f>", xml)
        self.assertIn("<m:rad>", xml)
        self.assertIn("<m:oMathPara", xml)
        self.assertNotIn("dfrac", xml)
        self.assertNotIn("begin{align", xml)

    def test_empty_text_is_rejected(self):
        with self.assertRaisesRegex(ValueError, "粘贴的文本不能为空"):
            convert("  ", str(self.work_dir / "empty.docx"))

    def test_blank_lines_use_plain_underscore_text_outside_math(self):
        output = self.work_dir / "blank-lines.docx"
        convert(
            "9. 单项式的系数是＿＿＿＿，次数是＿＿＿＿。\n"
            "10. y^4\\cdot y^5=________\n"
            "11. (-2)^0=_ _ _ _ _ _ _ _",
            str(output),
        )

        document = Document(output)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertNotIn("＿", text)
        self.assertIn("____", text)
        with zipfile.ZipFile(output) as package:
            xml = package.read("word/document.xml").decode("utf-8")
        self.assertNotIn('<w:u w:val="single"/>', xml)
        self.assertGreaterEqual(xml.count("________"), 1)
        self.assertNotIn("<m:sSub>", xml)

    def test_plain_text_with_unwrapped_latex_keeps_formula_structure(self):
        source = (
            "整式综合单元测试卷\n\n"
            "一、选择题（每题3分，共24分）\n"
            "1. 下列属于单项式的是（　　）\n"
            "A.x-y　B.\\dfrac{2}{5}x　C.x^2+3x　D.\\dfrac{1}{a}\n"
            "\\begin{align*}\n"
            "y^2&=x^2+1\\\\\n"
            "z&=\\sqrt{y}\n"
            "\\end{align*}"
        )
        output = self.work_dir / "plain-text-formula.docx"

        convert(source, str(output))

        document = Document(output)
        self.assertEqual(document.paragraphs[0].text, "整式综合单元测试卷")
        self.assertEqual(document.paragraphs[0].alignment, 1)
        with zipfile.ZipFile(output) as package:
            xml = package.read("word/document.xml").decode("utf-8")
        self.assertGreaterEqual(xml.count("<m:oMath"), 4)
        self.assertGreaterEqual(xml.count("<m:f>"), 2)
        self.assertIn("<m:sSup>", xml)
        self.assertIn("<m:rad>", xml)
        self.assertIn("<m:oMathPara", xml)
        self.assertNotIn("dfrac", xml)
        self.assertNotIn("begin{align", xml)

    def test_pdf_math_mode_embeds_formula_images_instead_of_omml(self):
        source = (
            "# 公式 PDF 测试\n\n"
            "1. 计算 $x^2 + \\dfrac{1}{2} + \\sqrt{y}$\n\n"
            "$$\\begin{align*}\n"
            "\\text{原式}&=x^2+4x+4\\\\\n"
            "&=\\boldsymbol{4x+13}\n"
            "\\end{align*}$$"
        )
        output = self.work_dir / "formula-images.docx"

        convert(source, str(output), math_mode="image")

        with zipfile.ZipFile(output) as package:
            xml = package.read("word/document.xml").decode("utf-8")
            media = [name for name in package.namelist() if name.startswith("word/media/")]
            image_bytes = [package.read(name) for name in media]
        self.assertNotIn("<m:oMath", xml)
        self.assertGreaterEqual(len(media), 2)
        self.assertTrue(
            all(Image.open(BytesIO(content)).convert("RGBA").getchannel("A").getbbox() for content in image_bytes)
        )

    def test_invalid_math_mode_is_rejected(self):
        with self.assertRaisesRegex(ValueError, "不支持的公式渲染模式"):
            convert("公式 $x^2$", str(self.work_dir / "invalid.docx"), math_mode="bad")

    def test_unicode_minus_does_not_split_unwrapped_bold_formula(self):
        output = self.work_dir / "unicode-minus.docx"

        convert(
            "19. x^2 − 5x+2x − 10=\\boldsymbol{x^2 − 3x − 10}\n"
            "20. \\boldsymbol{4a^2 − 9b^2}",
            str(output),
            math_mode="image",
        )

        with zipfile.ZipFile(output) as package:
            xml = package.read("word/document.xml").decode("utf-8")
            media = [name for name in package.namelist() if name.startswith("word/media/")]
        # Each answer line stays as one complete image formula.  Previously
        # U+2212 split the command into unmatched fragments and showed
        # ``\boldsymbol{`` in the generated PDF.
        self.assertEqual(xml.count("<w:drawing>"), 2)
        self.assertEqual(len(media), 2)


if __name__ == "__main__":
    unittest.main()
