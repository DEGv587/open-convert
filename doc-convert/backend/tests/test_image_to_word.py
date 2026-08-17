import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from PIL import Image

from converters.image_to_word import convert


def _ocr_data(*rows):
    data = {key: [] for key in ("text", "block_num", "par_num", "line_num", "left", "top")}
    for text, block, paragraph, line, left, top in rows:
        data["text"].append(text)
        data["block_num"].append(block)
        data["par_num"].append(paragraph)
        data["line_num"].append(line)
        data["left"].append(left)
        data["top"].append(top)
    return data


class ImageToWordTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.work_dir = Path(self.temp_dir.name)

    def tearDown(self):
        self.temp_dir.cleanup()

    def _image(self, name: str) -> Path:
        path = self.work_dir / name
        Image.new("RGB", (800, 1200), (240, 20, 20)).save(path, format="PNG")
        return path

    @patch("converters.image_to_word.pytesseract.get_languages", return_value=["chi_sim", "eng"])
    @patch(
        "converters.image_to_word.pytesseract.image_to_data",
        return_value=_ocr_data(
            ("你好", 1, 1, 1, 20, 20),
            ("世界", 1, 1, 1, 100, 20),
            ("Editable", 1, 2, 1, 20, 70),
        ),
    )
    def test_single_image_creates_editable_unicode_text(self, _image_to_data, _get_languages):
        source = self._image("中文页面.png")
        output = self.work_dir / "output.docx"
        progress = []

        convert(str(source), str(output), lambda pct, stage="": progress.append((pct, stage)))

        document = Document(output)
        self.assertEqual(len(document.inline_shapes), 0)
        self.assertIn("你好世界", document.paragraphs[0].text)
        self.assertIn("Editable", document.paragraphs[1].text)
        self.assertIn('w:eastAsia="宋体"', document.paragraphs[0].runs[0]._element.xml)
        self.assertEqual(document.core_properties.title, "中文页面")
        self.assertEqual(progress[-1], (100, "Word 文档已生成"))

    @patch("converters.image_to_word.pytesseract.get_languages", return_value=["chi_sim", "eng"])
    @patch(
        "converters.image_to_word.pytesseract.image_to_data",
        side_effect=[
            _ocr_data(("first", 1, 1, 1, 20, 20)),
            _ocr_data(("second", 1, 1, 1, 20, 20)),
        ],
    )
    def test_multiple_images_keep_order_and_page_break(self, _image_to_data, _get_languages):
        first = self._image("01.png")
        second = self._image("02.png")
        output = self.work_dir / "ordered.docx"

        convert([str(first), str(second)], str(output))

        document = Document(output)
        self.assertEqual(len(document.inline_shapes), 0)
        self.assertEqual(document.paragraphs[0].text, "first")
        self.assertEqual(document.paragraphs[1].text, "")
        self.assertTrue(document.paragraphs[1].runs[0]._element.xml.find("w:br") >= 0)
        self.assertEqual(document.paragraphs[2].text, "second")

    @patch("converters.image_to_word.pytesseract.get_languages", return_value=["eng"])
    def test_missing_chinese_language_pack_has_clear_error(self, _get_languages):
        source = self._image("page.png")

        with self.assertRaisesRegex(RuntimeError, "缺少 OCR 语言包：chi_sim"):
            convert(str(source), str(self.work_dir / "output.docx"))

    @patch("converters.image_to_word.pytesseract.get_languages", return_value=["chi_sim", "eng"])
    def test_invalid_image_has_clear_error(self, _get_languages):
        source = self.work_dir / "broken.png"
        source.write_text("not an image", encoding="utf-8")

        with self.assertRaisesRegex(ValueError, "无法读取图片：broken.png"):
            convert(str(source), str(self.work_dir / "output.docx"))


if __name__ == "__main__":
    unittest.main()
